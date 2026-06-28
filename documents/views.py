import io
from collections import Counter, defaultdict

from django.contrib.auth.decorators import login_required
from django.core.paginator import Paginator
from django.db.models import Exists, F, OuterRef
from django.http import Http404, HttpResponse
from django.shortcuts import get_object_or_404, render

from documents.diffing import diff_articles
from documents.models import Article, Document, Link, Redaction
from documents.refs import build_corpus_links
from search.services import search_in_document

PAGE_SIZE = 20


def _current_published_or_404(document):
    """Текущая опубликованная редакция акта или Http404.

    Единственная точка, где «нет видимого читателю текста» превращается в 404 —
    раньше тот же блок дублировался в каждой публичной вьюхе акта.
    """
    redaction = document.redactions.current_published().first()
    if redaction is None:
        raise Http404("Нет опубликованной редакции")
    return redaction


def _diff_pair_or_404(document, from_pk):
    """Пара (older, current) опубликованных редакций для diff читателя.

    current — текущая опубликованная; older — опубликованная редакция `from_pk`
    того же акта. 404, если older отсутствует либо уже совпадает с текущей.
    """
    current = _current_published_or_404(document)
    older = get_object_or_404(
        Redaction,
        pk=from_pk,
        document=document,
        review_status=Redaction.ReviewStatus.PUBLISHED,
    )
    if older.pk == current.pk:
        raise Http404("Редакция уже текущая — сравнивать не с чем")
    return older, current


def _build_article_tree(articles):
    """Проставить `child_nodes` каждой статье и вернуть корни (parent=None).

    Шаблон рендерит дерево рекурсивно; собираем связи одним проходом без
    доп. запросов к БД (articles уже выбраны списком).
    """
    children_map = defaultdict(list)
    for a in articles:
        children_map[a.parent_id].append(a)
    for a in articles:
        a.child_nodes = children_map[a.id]
    return children_map[None]


@login_required
def document_list(request):
    """Список актов с текущей опубликованной редакцией (с пагинацией)."""
    current = Redaction.objects.filter(document=OuterRef("pk")).current_published()
    documents = Document.objects.filter(Exists(current)).order_by("title")
    page_obj = Paginator(documents, PAGE_SIZE).get_page(request.GET.get("page"))

    # Слаги актов в избранном пользователя — для ★-переключателя в списке.
    # Ленивый импорт: documents не зависит от bookmarks на уровне модуля.
    from bookmarks.models import Bookmark

    bookmarked = set(
        Bookmark.objects.filter(user=request.user).values_list("document__slug", flat=True)
    )

    template = (
        "documents/_list_items.html"
        if request.headers.get("HX-Request")
        else "documents/document_list.html"
    )
    return render(request, template, {"page_obj": page_obj, "bookmarked": bookmarked})


@login_required
def document_detail(request, slug):
    """Страница-ридер акта: дерево статей, панели связей, кнопки AI."""
    document = get_object_or_404(Document, slug=slug)
    redaction = _current_published_or_404(document)

    articles = list(redaction.articles.all())
    article_tree = _build_article_tree(articles)
    kind_counts = Counter(a.kind for a in articles)
    # Якоря статей этого акта — для внутренних гиперссылок «ст. N» в тексте.
    anchors = {a.anchor for a in articles if a.anchor}
    visible_statuses = [Link.Status.CONFIRMED]
    if request.user.is_staff:
        visible_statuses.append(Link.Status.SUGGESTED)
    outgoing = document.outgoing_links.filter(status__in=visible_statuses).select_related(
        "to_document"
    )
    amendments = [
        link
        for link in outgoing
        if link.link_type in (Link.LinkType.AMENDS, Link.LinkType.AMENDED_BY)
    ]
    references = [link for link in outgoing if link.link_type == Link.LinkType.REFERENCES]
    incoming = document.incoming_links.filter(status__in=visible_statuses).select_related(
        "from_document"
    )
    published_redactions = document.redactions.published()

    return render(
        request,
        "documents/document_detail.html",
        {
            "document": document,
            "redaction": redaction,
            "article_tree": article_tree,
            "links": {"anchors": anchors, **build_corpus_links(exclude_slug=document.slug)},
            "show_ai": True,  # кнопка «разъяснить» — только в интерактивном reader, не в печати
            "amendments": amendments,
            "references": references,
            "incoming": incoming,
            "is_curator": request.user.is_staff,
            "published_redactions": published_redactions,
            "section_count": kind_counts.get(Article.Kind.SECTION, 0),
            "chapter_count": kind_counts.get(Article.Kind.CHAPTER, 0),
            "article_count": kind_counts.get(Article.Kind.ARTICLE, 0),
            "point_count": kind_counts.get(Article.Kind.POINT, 0),
            "appendix_count": kind_counts.get(Article.Kind.APPENDIX, 0),
        },
    )


@login_required
def article_explain(request, slug, anchor):
    """AI-разъяснение одной статьи «простыми словами» (htmx-ленивый партиал).

    Стоит 0 по умолчанию — вызывается только по кнопке у статьи. Без ключа/при
    ошибке API — режим `unavailable`, исходный текст статьи остаётся виден.
    """
    from assistant.article_explain import explain_article

    document = get_object_or_404(Document, slug=slug)
    redaction = _current_published_or_404(document)
    article = get_object_or_404(Article, redaction=redaction, anchor=anchor)
    explanation = explain_article(article.text)
    return render(
        request,
        "documents/_article_explanation.html",
        {"explanation": explanation},
    )


@login_required
def document_search(request, slug):
    """Поиск по тексту акта (HTMX-эндпоинт) — возвращает частичку с совпадениями."""
    document = get_object_or_404(Document, slug=slug)
    query = request.GET.get("q", "").strip()
    hits = search_in_document(document, query) if query else []
    return render(
        request,
        "documents/_find_results.html",
        {"document": document, "query": query, "hits": hits},
    )


@login_required
def changes_feed(request):
    """Лента изменений: недавно опубликованные редакции, новые сверху."""
    published = Redaction.objects.published()
    feed = published.select_related("document").order_by(
        F("published_at").desc(nulls_last=True), "-redaction_date"
    )
    page_obj = Paginator(feed, PAGE_SIZE).get_page(request.GET.get("page"))

    # Для «что изменилось» нужен pk предыдущей опубликованной редакции того же
    # документа. Один доп. запрос по документам страницы вместо N+1.
    doc_ids = {r.document_id for r in page_obj}
    history = defaultdict(list)  # doc_id -> [(redaction_date, pk), ...] по возрастанию
    for doc_id, red_date, pk in (
        published.filter(document_id__in=doc_ids)
        .order_by("redaction_date")
        .values_list("document_id", "redaction_date", "pk")
    ):
        history[doc_id].append((red_date, pk))
    for entry in page_obj:
        entry.prev_pk = next(
            (
                pk
                for red_date, pk in reversed(history[entry.document_id])
                if red_date < entry.redaction_date
            ),
            None,
        )

    return render(request, "documents/changes_feed.html", {"page_obj": page_obj})


@login_required
def redaction_diff(request, slug, from_pk):
    """Изменения «прошлая редакция → текущая» для читателя. Read-only."""
    document = get_object_or_404(Document, slug=slug)
    older, current = _diff_pair_or_404(document, from_pk)
    # diff_articles(база, новая): older — база, current — «новая»; имена параметров
    # функции (current_articles/draft_articles) — из admin-сценария, НЕ менять порядок.
    diffs = [
        d
        for d in diff_articles(list(older.articles.all()), list(current.articles.all()))
        if d.status != "same"
    ]
    return render(
        request,
        "documents/redaction_diff.html",
        {"document": document, "older": older, "current": current, "diffs": diffs},
    )


@login_required
def redaction_diff_explain(request, slug, from_pk):
    """AI-объяснение изменений «прошлая → текущая» (htmx-партиал, ленивая загрузка).

    По умолчанию ничего не стоит: вызывается только по кнопке на странице diff.
    Без ключа/при ошибке API — режим `unavailable`, сам diff остаётся виден.
    """
    from assistant.diff_explain import explain_diff

    document = get_object_or_404(Document, slug=slug)
    older, current = _diff_pair_or_404(document, from_pk)

    older_by_num = {a.number: a for a in older.articles.all()}
    newer_by_num = {a.number: a for a in current.articles.all()}
    changes = [
        {
            "number": d.number,
            "status": d.status,
            "old_text": getattr(older_by_num.get(d.number), "text", "") or "",
            "new_text": getattr(newer_by_num.get(d.number), "text", "") or "",
        }
        for d in diff_articles(list(older.articles.all()), list(current.articles.all()))
        if d.status != "same"
    ]
    explanation = explain_diff(changes)
    return render(
        request,
        "documents/_diff_explanation.html",
        {"explanation": explanation},
    )


@login_required
def document_print(request, slug):
    """Версия для печати: чистая standalone-страница с полным текстом акта."""
    document = get_object_or_404(Document, slug=slug)
    redaction = _current_published_or_404(document)

    articles = list(redaction.articles.all())
    article_tree = _build_article_tree(articles)
    # Якоря статей этого акта — _article_node.html линкует «ст. N» в тексте.
    anchors = {a.anchor for a in articles if a.anchor}

    return render(
        request,
        "documents/document_print.html",
        {
            "document": document,
            "redaction": redaction,
            "article_tree": article_tree,
            "links": {"anchors": anchors, **build_corpus_links(exclude_slug=document.slug)},
        },
    )


_DOCX_HEADING_LEVEL = {
    Article.Kind.SECTION: 1,
    Article.Kind.CHAPTER: 2,
    Article.Kind.APPENDIX: 2,
    Article.Kind.ARTICLE: 3,
    Article.Kind.POINT: 4,
}
_DOCX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


def _articles_in_reading_order(articles):
    """DFS дерева статей в порядке чтения (раздел→глава→статья, по `order`)."""
    children = defaultdict(list)
    for a in articles:
        children[a.parent_id].append(a)

    def walk(parent_id):
        for a in sorted(children[parent_id], key=lambda x: x.order):
            yield a
            yield from walk(a.id)

    return list(walk(None))


@login_required
def document_export_docx(request, slug):
    """Экспорт акта в .docx (как у классических СПС «сохранить в Word»).

    Форматирует уже опубликованный (кураторский) текст — не генерирует новых норм.
    """
    from docx import Document as DocxDocument

    document = get_object_or_404(Document, slug=slug)
    redaction = _current_published_or_404(document)

    docx = DocxDocument()
    docx.add_heading(document.title, level=0)
    meta = document.get_doc_type_display()
    if document.official_number:
        meta += f" № {document.official_number}"
    meta += f" · редакция от {redaction.redaction_date:%d.%m.%Y}"
    docx.add_paragraph(meta)
    docx.add_paragraph(
        "Не является источником официального опубликования и не заменяет "
        "юридическую консультацию. Официальный источник: pravo.gov.ru."
    )
    if document.source_url:
        docx.add_paragraph(f"Официальный первоисточник: {document.source_url}")

    for a in _articles_in_reading_order(list(redaction.articles.all())):
        heading = a.get_kind_display()
        if a.number:
            heading += f" {a.number}"
        if a.title:
            heading += f". {a.title}"
        docx.add_heading(heading, level=_DOCX_HEADING_LEVEL.get(a.kind, 3))
        if a.text:
            docx.add_paragraph(a.text)

    buf = io.BytesIO()
    docx.save(buf)
    response = HttpResponse(buf.getvalue(), content_type=_DOCX_CONTENT_TYPE)
    response["Content-Disposition"] = f'attachment; filename="{document.slug}.docx"'
    return response
